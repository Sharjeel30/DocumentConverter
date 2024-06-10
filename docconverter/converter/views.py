from django.shortcuts import render, redirect
from .forms import DocumentForm
from .models import Document
import os
from docx import Document as DocxDocument
from pdf2docx import parse
import markdown2

# Create your views here.

def handle_uploaded_file(file, format):
    input_path = file.path
    output_path = os.path.splitext(input_path)[0] + '.' + format

    if file.name.endswith('.docx'):
        if format == 'pdf':
            doc = DocxDocument(input_path)
            # Conversion logic to PDF
            doc.save(output_path)
        elif format == 'md':
            doc = DocxDocument(input_path)
            content = '\n'.join([p.text for p in doc.paragraphs])
            with open(output_path, 'w') as f:
                f.write(markdown2.markdown(content))
    elif file.name.endswith('.pdf'):
        if format == 'docx':
            parse(input_path, output_path)
        elif format == 'md':
            # Conversion logic to Markdown
            pass
    elif file.name.endswith('.md'):
        if format == 'docx':
            content = markdown2.markdown_path(input_path)
            doc = DocxDocument()
            doc.add_paragraph(content)
            doc.save(output_path)
        elif format == 'pdf':
            # Conversion logic to PDF
            pass

    return output_path

def upload_file(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save()
            format = request.POST.get('format')
            converted_path = handle_uploaded_file(document.original_file, format)
            document.converted_file = converted_path
            document.save()
            return redirect('converter:convert_success', pk=document.pk)
    else:
        form = DocumentForm()
    return render(request, 'converter/upload.html', {'form': form})

def convert_success(request, pk):
    document = Document.objects.get(pk=pk)
    return render(request, 'converter/success.html', {'document': document})